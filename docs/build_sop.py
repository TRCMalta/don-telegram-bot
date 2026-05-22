"""Generate Don CoS SOP — Word document for governance + EU AI Act audit.

Output:
  ~/Documents/Obsidian/Claude/Reference/Don-CoS-SOP-v1.0.docx
  ~/Documents/Claude/don-telegram-bot/docs/SOP-v1.0.docx
"""

from __future__ import annotations

import datetime as dt
import pathlib

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


DOC_VERSION = "1.0"
DOC_DATE = dt.date.today().isoformat()
OWNER = "Jonathan Cremona (Founder, The Remarkable Collective)"
APPROVER = "Jonathan Cremona"

PRIMARY = RGBColor(0x0A, 0x0C, 0x14)   # ink
ACCENT  = RGBColor(0xFF, 0x5C, 0x41)   # coral
GREY    = RGBColor(0x4A, 0x4F, 0x5C)


# ---------------------------------------------------------------------------
# Helpers (identical to Barry template)
# ---------------------------------------------------------------------------

def add_heading(doc, text, level, color=None):
    h = doc.add_heading(text, level=level)
    if color is not None:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Cm(0.63 * (level + 1))
    return p


def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = str(val)
            for p in t.rows[ri].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def page_break(doc):
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

doc = Document()

# Page margins
sec = doc.sections[0]
sec.top_margin    = Cm(2.2)
sec.bottom_margin = Cm(2.2)
sec.left_margin   = Cm(2.4)
sec.right_margin  = Cm(2.4)

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


# =========================================================================
# Cover
# =========================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run("STANDARD OPERATING PROCEDURE")
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = ACCENT

sub = doc.add_paragraph()
r = sub.add_run("Don — AI Chief of Staff (Telegram)")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = PRIMARY

p = doc.add_paragraph()
r = p.add_run(
    "Conversational AI assistant deployed for Mike Roberts (CEO, The Remarkable Collective). "
    "Provides real-time access to Microsoft 365 mail and calendar, Think Talent Odoo CRM, "
    "web research, voice message processing, and draft-and-confirm business communications, "
    "operating exclusively via Telegram with an explicit human-approval gate on all mutating actions."
)
r.font.size = Pt(12)
r.font.color.rgb = GREY

doc.add_paragraph()  # spacer
add_table(
    doc,
    ["Field", "Value"],
    [
        ["Document title",    "Don CoS SOP"],
        ["Version",           DOC_VERSION],
        ["Effective date",    DOC_DATE],
        ["Owner",             OWNER],
        ["Approver",          APPROVER],
        ["Review cycle",      "Quarterly or on material change"],
        ["Next review due",   (dt.date.today() + dt.timedelta(days=90)).isoformat()],
        ["Classification",    "Internal — Confidential"],
        ["Regulatory framing","EU AI Act (Reg. 2024/1689), GDPR (Reg. 2016/679), Malta DPA"],
    ],
    col_widths=[5.5, 11],
)

page_break(doc)


# =========================================================================
# 1. Purpose and scope
# =========================================================================
add_heading(doc, "1. Purpose and scope", 1, PRIMARY)
add_para(
    doc,
    "This Standard Operating Procedure documents the design, operation, oversight, and lifecycle "
    "of \"Don\" -- an AI Chief of Staff deployed on Telegram for Mike Roberts, Group CEO of The "
    "Remarkable Collective (TRC). The document is intended for:",
)
add_bullet(doc, "Internal governance -- verifying the system operates within stated bounds.")
add_bullet(doc, "External audit -- demonstrating compliance with the EU AI Act and GDPR.")
add_bullet(doc, "Operational continuity -- enabling any future operator to run, maintain, or decommission Don safely.")

add_heading(doc, "1.1 Scope", 2, PRIMARY)
add_para(doc, "In scope:")
add_bullet(doc, "Conversational AI layer (Anthropic Claude Sonnet via Anthropic API).")
add_bullet(doc, "Voice message transcription (Groq Whisper API).")
add_bullet(doc, "Microsoft 365 integration: read Mike's mailbox and calendar; draft and send emails (confirmation-gated); create calendar events (confirmation-gated).")
add_bullet(doc, "Odoo CRM integration: read and update Think Talent pipeline, contacts, products, events, invoices, and sales orders; log notes and create/update leads (confirmation-gated).")
add_bullet(doc, "Web research: Brave Search API and arbitrary public URL fetching.")
add_bullet(doc, "Persistent memory: lightweight preference and fact store on Railway volume.")
add_bullet(doc, "Cloud hosting on Railway with a persistent volume.")

add_para(doc, "Out of scope:")
add_bullet(doc, "Ceek Talent recruitment pipeline. That surface is managed by a separate system (Sam).")
add_bullet(doc, "Publishing to any public channel (LinkedIn, social media, website). Don is private-channel only.")
add_bullet(doc, "Think and Consult sub-brand operations beyond what appears in the shared Odoo instance.")
add_bullet(doc, "Any Firefish (ATS) or WhatsApp integrations.")

page_break(doc)


# =========================================================================
# 2. System overview
# =========================================================================
add_heading(doc, "2. System overview", 1, PRIMARY)
add_para(
    doc,
    "Don is a single-tenant Node.js (ESM) service running 24/7 on Railway. It receives messages "
    "exclusively from a whitelisted set of Telegram user IDs (Mike Roberts and Jonathan Cremona). "
    "All write actions -- sending email, creating calendar events, logging CRM notes, creating or "
    "updating leads -- require explicit text confirmation from Mike before the relevant tool is called. "
    "Read-only queries run freely without a gate. No output ever reaches a public channel.",
)

add_heading(doc, "2.1 High-level architecture", 2, PRIMARY)
add_bullet(doc, "Telegram webhook delivers messages to the Express server hosted on Railway.")
add_bullet(doc, "Allowlist check: messages from non-whitelisted Telegram user IDs are silently dropped.")
add_bullet(doc, "Voice messages are downloaded from Telegram, transcribed via Groq Whisper (whisper-large-v3-turbo), and treated as text input.")
add_bullet(doc, "The text input is passed to Anthropic Claude Sonnet (claude-sonnet-4-20250514) with the full system prompt (knowledge base + Mike profile + memory block) and conversation history.")
add_bullet(doc, "Claude invokes tools as needed: Odoo XML-RPC queries, Microsoft Graph API calls, Brave web search, or public URL fetch.")
add_bullet(doc, "For mutating tools (send_mike_email, create_mike_calendar_event, log_odoo_note, create_odoo_lead, update_odoo_lead), Claude presents a draft in Telegram and tool execution is blocked until Mike confirms.")
add_bullet(doc, "After each reply, a background Haiku call (claude-haiku-4-20250514) extracts new learnings and appends them to don-memory.json on the Railway volume.")
add_bullet(doc, "Conversation history is in-process only (LRU, 20 turns max per chat, wiped on restart). Memory facts persist on the volume.")

add_heading(doc, "2.2 Decision authority", 2, PRIMARY)
add_para(
    doc,
    "Don has no autonomous authority over mutating actions. Every write to Microsoft 365 or Odoo "
    "requires explicit confirmation from Mike Roberts in the Telegram chat. Read-only data retrieval "
    "is autonomous. The system can be stopped at any time by removing the Telegram bot token, "
    "revoking the MS Graph client secret, or stopping the Railway service.",
)

page_break(doc)


# =========================================================================
# 3. EU AI Act classification
# =========================================================================
add_heading(doc, "3. EU AI Act classification", 1, PRIMARY)
add_para(
    doc,
    "Regulation (EU) 2024/1689 -- the Artificial Intelligence Act -- entered into force on "
    "1 August 2024. The classifications, roles, and obligations below apply to Don as deployed "
    "on the effective date of this SOP.",
)

add_heading(doc, "3.1 Role under the Act", 2, PRIMARY)
add_table(
    doc,
    ["Role", "Status", "Reasoning"],
    [
        ["Provider (Art. 3(3))", "No",
         "Don does not develop or place a general-purpose or high-risk AI model on the EU market. "
         "It integrates models supplied by Anthropic (Claude) and Groq (Whisper)."],
        ["Deployer (Art. 3(4))", "YES",
         "Jonathan Cremona, sole proprietor of The Remarkable Collective, deploys Don under his "
         "professional authority for use by a named individual (Mike Roberts, CEO)."],
        ["Importer / Distributor", "No",
         "No third-party AI system is placed on the EU market by this deployment."],
        ["Affected persons", "Identified",
         "Primary affected natural person is Mike Roberts, who voluntarily uses the system. "
         "CRM queries may incidentally retrieve data relating to TRC clients and prospects."],
    ],
    col_widths=[3.5, 1.5, 11.5],
)

add_heading(doc, "3.2 Risk classification", 2, PRIMARY)
add_para(doc, "Don is classified as LIMITED RISK under the EU AI Act.", bold=True)
add_bullet(
    doc,
    "It is NOT a prohibited practice under Article 5. Don does not engage in subliminal manipulation, "
    "social scoring, biometric categorisation, real-time remote biometric identification, or any "
    "other practice listed in Art. 5.",
)
add_bullet(
    doc,
    "It is NOT a high-risk AI system under Article 6 / Annex III. An internal executive assistant "
    "operating on a private Telegram channel does not fall within any of the eight Annex III categories "
    "(biometrics, critical infrastructure, education, employment selection, essential services, law "
    "enforcement, migration, administration of justice).",
)
add_bullet(
    doc,
    "It does interact with a natural person (Mike Roberts) in a conversational context, so the "
    "transparency obligations of Article 50 are considered below.",
)
add_bullet(
    doc,
    "It uses general-purpose AI models (Claude, Groq Whisper) as components. The Article 53-55 "
    "obligations on GPAI providers fall on Anthropic and Groq respectively, not on the deployer.",
)

add_heading(doc, "3.3 Article 50 transparency obligations", 2, PRIMARY)
add_para(
    doc,
    "Article 50(1) requires deployers of AI systems that interact with natural persons to ensure "
    "those persons are informed they are interacting with an AI system, unless this is obvious from "
    "context. Article 50(4) requires disclosure for AI-generated text intended to inform the public "
    "on matters of public interest (editorial-control exemption available).",
)
add_para(
    doc,
    "Position taken by this SOP: Mike Roberts is fully aware that Don is an AI system -- the system "
    "introduced itself on first contact and the Telegram bot name and handle make this explicit. "
    "No output from Don is directed at the public. The Art. 50(1) requirement is met by design. "
    "Art. 50(4) does not apply because Don's output is never published to a public channel.",
    italic=True,
)

add_heading(doc, "3.4 GPAI dependency declaration", 2, PRIMARY)
add_table(
    doc,
    ["Model", "Provider", "Use", "EU AI Act tier"],
    [
        ["claude-sonnet-4-20250514", "Anthropic PBC (US)",
         "Primary conversation, tool-use orchestration, draft generation",
         "GPAI without systemic risk (provider obligation)"],
        ["claude-haiku-4-20250514", "Anthropic PBC (US)",
         "Background memory extraction (fire-and-forget after each reply)",
         "GPAI without systemic risk (provider obligation)"],
        ["whisper-large-v3-turbo", "Groq, Inc. (US) -- OpenAI Whisper model",
         "Voice message transcription",
         "GPAI without systemic risk (provider obligation)"],
    ],
    col_widths=[5.0, 3.0, 5.5, 4.0],
)

page_break(doc)


# =========================================================================
# 4. Connected services and data flows
# =========================================================================
add_heading(doc, "4. Connected services and data flows", 1, PRIMARY)
add_para(doc, "Every external service Don connects to is listed below.")

add_table(
    doc,
    ["Service", "Vendor", "Purpose", "Data sent", "Region", "Legal basis"],
    [
        ["Anthropic Claude API", "Anthropic PBC (US)",
         "Conversation, tool-use orchestration, memory extraction",
         "Message text, system prompt (knowledge base), tool results. No voice audio.",
         "US (EU SCCs)", "Legitimate interest"],
        ["Groq API (Whisper)", "Groq, Inc. (US)",
         "Voice message transcription",
         "OGG audio file of Mike's voice message. Not stored by Groq beyond transcription.",
         "US (EU SCCs)", "Legitimate interest"],
        ["Brave Search API", "Brave Software, Inc. (US)",
         "Web search queries on Mike's behalf",
         "Search query string. No personal data beyond query text.",
         "US", "Legitimate interest"],
        ["Telegram Bot API", "Telegram Messenger Inc. (International)",
         "Primary user-facing channel (messages in, replies out)",
         "Message text, voice files, callback data between Mike's phone and Don.",
         "International", "Legitimate interest (private communication)"],
        ["Microsoft Graph API", "Microsoft Ireland Operations Ltd (EU)",
         "Read/write Mike's M365 mailbox and calendar",
         "Email metadata and bodies, calendar events, attendee addresses. OAuth 2.0 token.",
         "EU (Azure Ireland)", "Legitimate interest + contract (M365 ToS)"],
        ["Odoo XML-RPC", "Odoo SA / thinktalent.com.mt (Malta/EU)",
         "Read/update Think Talent CRM pipeline, contacts, invoices",
         "CRM records including client/prospect names, company names, deal values. API key.",
         "Malta (EU)", "Legitimate interest (own business data)"],
        ["Public URLs (fetch_webpage)", "Various (public web)",
         "Full-text content retrieval for research queries",
         "HTTP GET to public URLs. No credentials or personal data sent.",
         "Various", "Legitimate interest"],
        ["Railway", "Railway Corp. (US)",
         "Cloud hosting + persistent volume (don-memory.json)",
         "Service code, environment variables (secrets), runtime logs, memory file.",
         "US (or selected region)", "Legitimate interest"],
        ["GitHub", "GitHub, Inc. (US)",
         "Source code repository (public)",
         "Code only. No secrets, no personal data, no memory state.",
         "US (EU SCCs)", "Legitimate interest"],
    ],
    col_widths=[3.0, 3.0, 3.5, 4.5, 2.0, 2.5],
)

add_heading(doc, "4.1 Sub-processor list", 2, PRIMARY)
add_bullet(doc, "Anthropic PBC -- model inference. DPA available at anthropic.com/legal.")
add_bullet(doc, "Groq, Inc. -- voice transcription. Groq privacy policy and DPA apply.")
add_bullet(doc, "Brave Software, Inc. -- web search. Brave Search API terms apply.")
add_bullet(doc, "Microsoft Ireland Operations Ltd -- M365 email and calendar. Microsoft DPA (EU) applies.")
add_bullet(doc, "Odoo SA / thinktalent.com.mt -- CRM data. Odoo.com DPA applies.")
add_bullet(doc, "Telegram Messenger Inc. -- message delivery channel.")
add_bullet(doc, "Railway Corp. -- compute and volume hosting.")
add_bullet(doc, "GitHub, Inc. -- code hosting.")

add_heading(doc, "4.2 Data flow diagram (textual)", 2, PRIMARY)
add_para(
    doc,
    "Mike's Telegram message (text or voice)  -->  Telegram Bot API  -->  Railway Express webhook  -->  "
    "Allowlist check (drop if not whitelisted)  -->  [Voice: Groq Whisper transcription]  -->  "
    "Claude Sonnet (system prompt + history + memory)  -->  [Tool calls as needed: "
    "MS Graph / Odoo XML-RPC / Brave Search / fetch URL]  -->  Claude reply text  -->  "
    "Telegram Bot API  -->  Mike's phone. "
    "Background: Claude Haiku extracts memory  -->  don-memory.json on Railway /data volume.",
)

page_break(doc)


# =========================================================================
# 5. Account isolation and security controls
# =========================================================================
add_heading(doc, "5. Account isolation and security controls", 1, PRIMARY)

add_heading(doc, "5.1 Access control -- Telegram allowlist", 2, PRIMARY)
add_para(
    doc,
    "Don inspects the Telegram user ID of every incoming message. Only user IDs listed in the "
    "ALLOWED_TELEGRAM_USERS environment variable are processed. Any message from an unlisted user "
    "is silently dropped before any processing occurs. This is enforced in the webhook handler before "
    "any API call is made.",
)

add_heading(doc, "5.2 Microsoft Graph -- least-privilege scopes", 2, PRIMARY)
add_para(
    doc,
    "The Azure App Registration (\"Don TRC Bot\") is granted only the scopes required: Mail.Read, "
    "Mail.Send, Calendars.Read, Calendars.ReadWrite for Mike's mailbox only (MIKE_EMAIL). "
    "The client secret (MS_CLIENT_SECRET) lives exclusively in Railway environment variables and is "
    "never written to the GitHub repository, Obsidian vault, or any chat transcript.",
)

add_heading(doc, "5.3 Secret handling", 2, PRIMARY)
add_bullet(doc, "All credentials live in Railway environment variables, never in the GitHub repository (which is public).")
add_bullet(doc, "Secrets: TELEGRAM_TOKEN, ANTHROPIC_API_KEY, MS_TENANT_ID, MS_CLIENT_ID, MS_CLIENT_SECRET, ODOO_API_KEY, GROQ_API_KEY, BRAVE_API_KEY, MIKE_EMAIL, ALLOWED_TELEGRAM_USERS.")
add_bullet(doc, "MS_CLIENT_SECRET rotation cadence: TBC -- owner to confirm. Recommended: annually or on any suspected compromise.")
add_bullet(doc, "No secret is ever logged. Application logs do not contain credential values.")
add_bullet(doc, "Local .env file (development only) is excluded from git via .gitignore.")

add_heading(doc, "5.4 Network exposure", 2, PRIMARY)
add_bullet(doc, "Public surface: one HTTPS endpoint (POST /webhook) receives Telegram updates. The path includes the bot token as a path component (Telegram's default webhook authentication mechanism).")
add_bullet(doc, "Additional read-only diagnostic endpoints: GET /health, GET /memory, GET /debug, GET /webhook-info.")
add_bullet(doc, "No unauthenticated write surface is exposed. POST /webhook validates via Telegram's token-in-path mechanism.")
add_bullet(doc, "Memory clear endpoint (DELETE /memory) requires direct Railway access or knowledge of the private URL -- no authentication token beyond Railway network controls.")
add_bullet(doc, "Live URL: https://don-telegram-bot-production.up.railway.app")

page_break(doc)


# =========================================================================
# 6. Human oversight
# =========================================================================
add_heading(doc, "6. Human oversight (EU AI Act Article 14 -- voluntary alignment)", 1, PRIMARY)
add_para(
    doc,
    "Although Don is not a high-risk AI system and Article 14 is not strictly mandatory, "
    "the controls below align with its principles.",
)
add_bullet(doc, "Confirmation gate for all mutating actions. Don presents a draft (email body, calendar event, CRM note, lead details) and waits for explicit approval before calling the write tool. Accepted confirmations: yes, send it, go ahead, do it, approved, confirmed, or a thumbs-up emoji.")
add_bullet(doc, "Read-only tools run freely. Queries to Odoo, M365 mailbox, M365 calendar, web search, and URL fetch do not require a confirmation gate.")
add_bullet(doc, "Voice input handled transparently. Voice messages are transcribed and the transcript is shown to Mike in the reply so he can correct any transcription error before acting on it.")
add_bullet(doc, "Memory is additive only. Don never deletes information Mike has provided; the memory file is human-readable JSON on the Railway volume and can be inspected or wiped via DELETE /memory.")
add_bullet(doc, "Kill switches: (a) remove TELEGRAM_TOKEN from Railway env to stop all inbound messages; (b) revoke MS_CLIENT_SECRET in Azure to disable M365 access; (c) revoke ODOO_API_KEY in Odoo to disable CRM access; (d) stop the Railway service to halt everything immediately.")
add_bullet(doc, "Max tool iterations: Claude is limited to 10 tool calls per conversational turn. If this limit is hit, Claude is forced to summarise what it found and return a text reply without further tool calls.")

page_break(doc)


# =========================================================================
# 7. Content / output controls
# =========================================================================
add_heading(doc, "7. Content and output controls", 1, PRIMARY)
add_para(
    doc,
    "Don's system prompt encodes explicit voice and behaviour rules enforced on every response.",
)
add_bullet(doc, "British English throughout. Colloquial register -- warm, efficient, never corporate.")
add_bullet(doc, "No em-dashes, en-dashes, or any Unicode dash variant in output. Hyphens in compound words are permitted.")
add_bullet(doc, "Lead with the answer. No preamble, no context-setting, no pleasantries.")
add_bullet(doc, "Banned words and phrases: Certainly, Absolutely, Of course, Great question, I'd be happy to, No problem, Just to let you know, leverage, synergies, holistic, game-changer, dive into, unlock, empower, robust, seamless, cutting-edge.")
add_bullet(doc, "Don never claims to be a human or denies being an AI if sincerely asked.")
add_bullet(doc, "Don never acts on a mutating request without a confirmation gate, regardless of how the instruction is phrased.")
add_bullet(doc, "Don's loyalty is to Mike (primary user) and Jonathan (owner/admin). It does not take instructions from other parties in the Telegram chat if the chat were ever expanded.")

page_break(doc)


# =========================================================================
# 8. GDPR position
# =========================================================================
add_heading(doc, "8. GDPR position", 1, PRIMARY)
add_para(doc, "Personal data processed by or through Don:")
add_bullet(doc, "Mike Roberts' email content -- names, addresses, message bodies, subject lines of emails in his M365 mailbox.")
add_bullet(doc, "Mike Roberts' calendar data -- meeting titles, times, attendee names and email addresses.")
add_bullet(doc, "Think Talent CRM data -- names, company names, email addresses, phone numbers, deal values, and other personal data of TRC clients and prospects stored in Odoo.")
add_bullet(doc, "Mike Roberts' voice messages -- OGG audio sent to Groq for transcription; not stored by Don or Groq beyond the transcription request.")
add_bullet(doc, "Persistent memory file (don-memory.json) -- Mike's expressed preferences, key facts about deals and projects, and notes on named individuals. Stored on Railway volume.")
add_bullet(doc, "Telegram message metadata -- user IDs, message timestamps, chat IDs. Processed by Telegram's infrastructure.")

add_heading(doc, "8.1 Controller identification", 2, PRIMARY)
add_para(
    doc,
    "The data controller for all processing described above is Jonathan Cremona, sole proprietor of "
    "The Remarkable Collective (Malta). Mike Roberts processes data in his capacity as an employee "
    "(CEO) of TRC, acting within the scope of his role. TRC is the controller for Odoo CRM data "
    "relating to clients and prospects.",
)

add_heading(doc, "8.2 GDPR articles addressed", 2, PRIMARY)
add_table(
    doc,
    ["Article", "Topic", "Status / Mitigation"],
    [
        ["Art. 5", "Data minimisation, purpose limitation, storage limitation",
         "Applied. Conversation history wiped on restart. Voice audio not stored. Memory file retains only summaries, not raw message content."],
        ["Art. 6", "Lawful basis",
         "Legitimate interest (TRC business operations) for email/calendar/CRM access. "
         "Implied consent for Mike's voluntary use of the system."],
        ["Art. 9", "Special categories of data",
         "Not intentionally processed. If special-category data appears in emails or CRM records, it is accessed only as part of Mike's legitimate business operations and is not extracted or stored separately."],
        ["Art. 13/14", "Information notice to data subjects",
         "TBC -- owner to confirm. Third parties whose data appears in Mike's inbox or CRM have not been specifically notified that an AI system may process their data. Include in TRC privacy notice update."],
        ["Art. 22", "Automated decision-making with legal or significant effect",
         "Not applicable. All consequential actions (email sends, calendar creates, CRM updates) require Mike's explicit confirmation. No solely automated decisions with legal effect are made."],
        ["Art. 25", "Data protection by design and by default",
         "Allowlist restricts access to named users. Secrets in Railway env only. Conversation history not persisted. Minimal memory extraction."],
        ["Art. 30", "Records of processing activities (RoPA)",
         "This SOP serves as the RoPA entry for Don. TRC must include this processing activity in its full RoPA."],
        ["Art. 32", "Security of processing",
         "TLS in transit for all API calls. Secrets in Railway env only. No public write surface. Railway volume access requires Railway credentials."],
        ["Art. 33/34", "Breach notification",
         "In event of suspected credential compromise or data leak: (a) revoke affected credential immediately; (b) notify IDPC within 72 hours if personal data of data subjects (not just Mike) is likely affected; (c) document in SOP change history."],
        ["Art. 35", "DPIA threshold",
         "DPIA not triggered for the current deployment scope. If Don's access is expanded to process employee data at scale or monitor individuals systematically, a DPIA must be conducted first."],
        ["Art. 44-49", "International transfers",
         "US sub-processors: Anthropic (SCCs), Groq (SCCs -- TBC owner to confirm DPA exists), Brave (SCCs -- TBC owner to confirm), Railway (SCCs -- TBC owner to confirm). Microsoft is EU-based. Odoo is Malta (EU)."],
    ],
    col_widths=[2.5, 5.0, 10.0],
)

page_break(doc)


# =========================================================================
# 9. Operational procedures
# =========================================================================
add_heading(doc, "9. Operational procedures", 1, PRIMARY)

add_heading(doc, "9.1 Start / stop / redeploy", 2, PRIMARY)
add_para(doc, "Railway is the production runtime. All CLI commands assume railway CLI installed and authenticated.")
add_bullet(doc, "Start (idempotent): railway up --detach")
add_bullet(doc, "Stop: from Railway dashboard, click Stop on the don-telegram-bot service. Or: railway service don-telegram-bot && railway down")
add_bullet(doc, "Redeploy after code change: git push origin main -- Railway auto-deploys on push to main.")
add_bullet(doc, "View live logs: railway logs --deployment")
add_bullet(doc, "Update environment variable: railway variables --set KEY=VALUE")
add_bullet(doc, "Set Telegram webhook after redeploy (if URL changes): GET https://don-telegram-bot-production.up.railway.app/set-webhook")

add_heading(doc, "9.2 Credential rotation", 2, PRIMARY)
add_bullet(doc, "MS_CLIENT_SECRET (Azure): generate new secret in Azure AD app registrations; update Railway env var; redeploy; verify via /debug endpoint.")
add_bullet(doc, "ODOO_API_KEY: generate new key in Odoo Settings > Users > API Keys; update Railway env var; redeploy.")
add_bullet(doc, "TELEGRAM_TOKEN: if bot is compromised, use BotFather /revoke; generate new token; update Railway env var; re-set webhook; redeploy.")
add_bullet(doc, "ANTHROPIC_API_KEY, GROQ_API_KEY, BRAVE_API_KEY: rotate in respective provider dashboards; update Railway env var; redeploy.")
add_bullet(doc, "After any rotation: verify via /debug endpoint and send a test message in Telegram.")
add_bullet(doc, "Log rotation date in Section 13 (Document history).")

add_heading(doc, "9.3 Updating Don's knowledge base", 2, PRIMARY)
add_bullet(doc, "Edit the buildSystemPrompt() function in server.js.")
add_bullet(doc, "Run: node --check server.js to verify syntax.")
add_bullet(doc, "Commit and push to main: railway auto-deploys.")
add_bullet(doc, "Send a test message in Telegram to confirm the new knowledge is accessible.")

add_heading(doc, "9.4 Memory management", 2, PRIMARY)
add_bullet(doc, "View current memory: GET https://don-telegram-bot-production.up.railway.app/memory")
add_bullet(doc, "Clear all memory: DELETE https://don-telegram-bot-production.up.railway.app/memory (use with caution -- irreversible)")
add_bullet(doc, "Memory file location on volume: /data/don-memory.json")
add_bullet(doc, "Memory is backed by the Railway volume -- survives redeployments but not volume deletion.")

add_heading(doc, "9.5 Incident response runbook", 2, PRIMARY)
add_bullet(doc, "Symptom: Don sends an email without Mike's confirmation. Action: (a) identify sent email in Mike's Sent Items; (b) recall if possible via M365; (c) revoke MS_CLIENT_SECRET immediately; (d) audit server logs for how the gate was bypassed; (e) document in change history and notify Jonathan.")
add_bullet(doc, "Symptom: Unknown user receiving Don responses. Action: check ALLOWED_TELEGRAM_USERS env var; revoke TELEGRAM_TOKEN immediately if compromised; audit webhook logs.")
add_bullet(doc, "Symptom: Groq transcription returns garbage. Action: Mike can type the message instead. Log the failure. Check GROQ_API_KEY validity.")
add_bullet(doc, "Symptom: Odoo queries return errors. Action: check ODOO_API_KEY validity in /debug; verify thinktalent.com.mt is reachable; check Odoo server status.")
add_bullet(doc, "Symptom: MS Graph returns 401. Action: MS_CLIENT_SECRET may have expired or been revoked. Run credential rotation procedure (9.2).")
add_bullet(doc, "Symptom: Railway service crashed. Action: railway redeploy --yes; tail logs for root cause.")

page_break(doc)


# =========================================================================
# 10. Audit log and retention
# =========================================================================
add_heading(doc, "10. Audit log and retention", 1, PRIMARY)
add_table(
    doc,
    ["Record", "Location", "Retained", "Purpose"],
    [
        ["Railway application logs",
         "Railway log viewer (deployment-scoped)",
         "30 days (Railway default)",
         "Operational debugging, error tracing, audit of tool calls and errors"],
        ["don-memory.json (preferences, key facts, people notes)",
         "Railway /data volume",
         "Until manually cleared or volume deleted",
         "Persistent context for Mike's preferences and key business facts"],
        ["Telegram message history",
         "Telegram cloud (Mike's device and Telegram servers)",
         "Telegram platform default",
         "Evidence of human approval for any mutating action"],
        ["Conversation history (in-process)",
         "Node.js process memory only",
         "Lost on restart (by design)",
         "Context window for the current session only"],
        ["Voice message audio",
         "Not stored -- downloaded from Telegram, sent to Groq, discarded",
         "Not retained by Don",
         "Transcription only; Groq processes and discards per their policy"],
        ["MS Graph email and calendar data",
         "Microsoft 365 (Mike's mailbox and calendar)",
         "Microsoft retention policy (not controlled by Don)",
         "Source of record for emails and meetings"],
        ["Odoo CRM records",
         "thinktalent.com.mt Odoo instance",
         "Odoo retention policy (not controlled by Don)",
         "Source of record for pipeline and contacts"],
        ["This SOP",
         "Obsidian vault + GitHub docs/",
         "Indefinite (versioned)",
         "Governance artefact"],
    ],
    col_widths=[4.0, 4.0, 3.5, 5.0],
)

page_break(doc)


# =========================================================================
# 11. Risk register
# =========================================================================
add_heading(doc, "11. Risk register", 1, PRIMARY)
add_table(
    doc,
    ["#", "Risk", "Likelihood", "Impact", "Mitigation", "Owner"],
    [
        ["R1", "Don sends an email without Mike's explicit confirmation",
         "Very low", "Reputational / operational damage",
         "Confirmation gate enforced in system prompt and tool dispatcher. Gate requires positive text confirmation. No tool call made without it.",
         "Jonathan"],
        ["R2", "MS_CLIENT_SECRET leaked via logs or repo",
         "Low", "Full access to Mike's M365 mailbox and calendar",
         "Secret in Railway env only. Never logged. GitHub repo is public -- secret must never enter the codebase.",
         "Jonathan"],
        ["R3", "Telegram bot token compromised -- third party receives Don's responses",
         "Low", "Privacy breach, potential data leak",
         "ALLOWED_TELEGRAM_USERS allowlist means responses only go to whitelisted IDs. Token rotation procedure documented.",
         "Jonathan"],
        ["R4", "Odoo API key compromised",
         "Low", "Unauthorised read/write of Think Talent CRM",
         "Key in Railway env only. Rotation procedure in 9.2.",
         "Jonathan"],
        ["R5", "Don provides incorrect legal or HR advice",
         "Medium", "Business risk if acted upon without verification",
         "System prompt frames all legal/HR content as reference material, not professional advice. Mike expected to verify with qualified advisers.",
         "Mike / Jonathan"],
        ["R6", "Third-party personal data in CRM processed without adequate GDPR notice",
         "Medium", "IDPC enforcement action",
         "Include Don as processing activity in TRC privacy notice and RoPA. Art. 13/14 action: TBC -- owner to confirm.",
         "Jonathan"],
        ["R7", "Voice transcription error causes Mike to act on incorrect information",
         "Low", "Operational mistake",
         "Transcript shown in reply before any action is taken. Mike can correct before confirming.",
         "Mike"],
        ["R8", "Railway volume data loss (don-memory.json)",
         "Very low", "Loss of accumulated preferences and facts",
         "System functions without memory (falls back to stateless mode). Memory is supplementary, not critical-path.",
         "Jonathan"],
        ["R9", "Anthropic or Groq API outage during use",
         "Low", "Don temporarily unavailable",
         "No fallback AI -- Mike must retry or use direct tools. Acceptable degradation for a non-critical-path tool.",
         "Jonathan"],
        ["R10", "EU AI Act classification challenged (Limited Risk upgraded to High Risk)",
         "Very low", "Compliance remediation required",
         "Annual review of classification against updated Annex III guidance. If employment-related decision support is added (e.g. HR selection queries), classification must be revisited.",
         "Jonathan"],
    ],
    col_widths=[1.0, 4.5, 2.0, 2.0, 5.5, 2.0],
)

page_break(doc)


# =========================================================================
# 12. Decommissioning
# =========================================================================
add_heading(doc, "12. Decommissioning procedure", 1, PRIMARY)
add_bullet(doc, "Step 1: notify Mike Roberts that Don will be stopped and confirm a handover date.")
add_bullet(doc, "Step 2: revoke Telegram bot token via BotFather (/revoke or /deletebot).")
add_bullet(doc, "Step 3: revoke MS_CLIENT_SECRET in Azure AD app registrations (\"Don TRC Bot\").")
add_bullet(doc, "Step 4: revoke ODOO_API_KEY in Odoo Settings.")
add_bullet(doc, "Step 5: remove all Railway environment variables to prevent accidental restart with live credentials.")
add_bullet(doc, "Step 6: export don-memory.json from the Railway volume to long-term archive (~/Documents/Obsidian/08-Archive/Don-CoS/).")
add_bullet(doc, "Step 7: archive the Railway service (do not delete -- preserves 30 days of logs for audit).")
add_bullet(doc, "Step 8: mark project as archived in 01-Projects/Don-CoS.md and move to 08-Archive/.")
add_bullet(doc, "Step 9: update this SOP -- add decommissioning date to Section 13. Store final version alongside archived memory.")

page_break(doc)


# =========================================================================
# 13. Document history
# =========================================================================
add_heading(doc, "13. Document history", 1, PRIMARY)
add_table(
    doc,
    ["Version", "Date", "Author", "Change"],
    [
        ["1.0", DOC_DATE, "Jonathan Cremona + Claude (Anthropic)",
         "Initial SOP. Covers Don v1.0 as deployed on Railway 2026-05-22. "
         "Includes MS Graph, Odoo, Groq voice, Brave Search, Malta employment law, EU GDPR knowledge."],
    ],
    col_widths=[2.0, 3.0, 6.5, 6.0],
)


# =========================================================================
# 14. Acceptance and signature
# =========================================================================
add_heading(doc, "14. Acceptance and signature", 1, PRIMARY)
add_para(
    doc,
    "By signing below, the deployer confirms the system operates within the bounds described, "
    "the EU AI Act classification is supported by the operational facts, and the GDPR position "
    "is accurate.",
)
doc.add_paragraph()
doc.add_paragraph()
add_para(doc, "Signature: ___________________________________________", size=12)
add_para(doc, "Name: Jonathan Cremona", size=12)
add_para(doc, "Role: Deployer (sole proprietor, The Remarkable Collective)", size=12)
add_para(doc, "Date: ____________________", size=12)


# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
out1 = pathlib.Path.home() / "Documents" / "Obsidian" / "Claude" / "Reference" / "Don-CoS-SOP-v1.0.docx"
out2 = pathlib.Path(__file__).parent / "SOP-v1.0.docx"
out1.parent.mkdir(parents=True, exist_ok=True)
out2.parent.mkdir(parents=True, exist_ok=True)
doc.save(out1)
doc.save(out2)
print(f"saved --> {out1}")
print(f"saved --> {out2}")
